#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для преобразования отчетов из Excel в Word
Анализирует файл esklp_full.xlsx и создает документы Word для каждого листа
"""

import os
import sys
import logging
from pathlib import Path
from datetime import datetime

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('excel_to_word.log', encoding='utf-8', mode='w'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


def set_cell_background(cell, color):
    """Установка цвета фона ячейки"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        from docx.oxml import OxmlElement
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), color)


def detect_header_row(df_raw, min_row=5):
    """
    Автоопределение строки заголовка
    Ищем первую строку после min_row, которая содержит осмысленные данные (не NaN и не Unnamed)
    
    Args:
        df_raw: DataFrame без заголовков
        min_row: Минимальный индекс строки для поиска заголовка
    
    Returns:
        Индекс строки заголовка
    """
    for idx in range(min_row, len(df_raw)):
        row = df_raw.iloc[idx]
        # Проверяем, есть ли в строке значения, которые не являются Unnamed
        non_unnamed_count = sum(1 for val in row if pd.notna(val) and not str(val).lower().startswith('unnamed'))
        if non_unnamed_count > 0:
            return idx
    return min_row


def load_sheet_with_header_detection(input_file, sheet_name):
    """
    Загрузка листа с автоопределением заголовка
    Для листа ЕСКЛП: заголовки в строке 6 (индекс 5), данные с 7 строки
    Для остальных листов: заголовки в строке 1 (индекс 0), данные со 2 строки
    
    Args:
        input_file: Путь к файлу Excel
        sheet_name: Имя листа
    
    Returns:
        DataFrame с правильными заголовками
    """
    # Для листа ЕСКЛП используем автоопределение заголовка
    if sheet_name == 'ЕСКЛП':
        # Сначала загружаем без заголовков для определения строки заголовка
        df_raw = pd.read_excel(input_file, sheet_name=sheet_name, header=None)
        
        # Определяем строку заголовка (начиная с минRow=5)
        header_row_idx = detect_header_row(df_raw, min_row=5)
        logger.info(f"  Заголовок найден в строке {header_row_idx + 1}")
        
        # Загружаем с правильной строкой заголовка
        df = pd.read_excel(input_file, sheet_name=sheet_name, header=header_row_idx)
    else:
        # Для остальных листов заголовок в первой строке
        df = pd.read_excel(input_file, sheet_name=sheet_name, header=0)
        logger.info(f"  Заголовок в строке 1")
    
    # Фильтруем колонки, исключая Unnamed
    valid_columns = [col for col in df.columns if not (pd.isna(col) or str(col).lower().startswith('unnamed'))]
    
    if len(valid_columns) < len(df.columns):
        logger.info(f"  Исключено {len(df.columns) - len(valid_columns)} колонок Unnamed")
    
    df = df[valid_columns]
    
    return df


def analyze_column_uniqueness(df, column):
    """
    Анализ уникальности значений в колонке
    
    Returns:
        dict: Статистика уникальности
    """
    values = df[column].dropna()
    total_count = len(values)
    unique_count = values.nunique()
    
    if total_count == 0:
        return {
            'total': 0,
            'unique': 0,
            'uniqueness_percent': 0,
            'is_unique': False,
            'duplicate_count': 0
        }
    
    uniqueness_percent = round(unique_count / total_count * 100, 2)
    duplicate_count = total_count - unique_count
    
    return {
        'total': total_count,
        'unique': unique_count,
        'uniqueness_percent': uniqueness_percent,
        'is_unique': uniqueness_percent == 100.0,
        'duplicate_count': duplicate_count
    }


def detect_pk_fk(df, column_name):
    """
    Автоопределение типа ключа (PK/FK) для колонки
    
    Returns:
        str: Тип ключа ('PK', 'FK', 'Potential PK', 'Potential FK', '')
    """
    col_lower = str(column_name).lower()
    uniqueness = analyze_column_uniqueness(df, column_name)
    
    # Проверка на первичный ключ по имени
    pk_patterns = ['id', 'guid', 'uuid', 'код', 'code', 'primary']
    is_pk_by_name = any(pattern in col_lower for pattern in pk_patterns)
    
    # Проверка на внешний ключ по имени
    fk_patterns = ['_id', 'id_', 'parent', 'родитель', 'foreign', 'ref', 'ссылка']
    is_fk_by_name = any(pattern in col_lower for pattern in fk_patterns)
    
    # Если колонка уникальна на 100% и имеет признаки PK
    if uniqueness['is_unique'] and is_pk_by_name:
        return 'PK'
    
    # Если колонка уникальна на 100% и похожа на идентификатор
    if uniqueness['is_unique'] and ('guid' in col_lower or 'uuid' in col_lower or col_lower == 'id'):
        return 'PK'
    
    # Если колонка имеет признаки FK
    if is_fk_by_name and uniqueness['uniqueness_percent'] < 100:
        return 'FK'
    
    # Потенциальный PK (уникален, но имя не явно указывает)
    if uniqueness['is_unique']:
        return 'Potential PK'
    
    # Потенциальный FK (имя указывает, но не уникален)
    if is_fk_by_name:
        return 'Potential FK'
    
    return ''


def create_word_report_from_excel(input_file=None, output_dir=None, mode='single'):
    """
    Создание отчетов Word из файла Excel
    
    Args:
        input_file: Путь к файлу Excel
        output_dir: Директория для сохранения отчетов Word
        mode: 'single' - все листы в одном документе, 'separate' - каждый лист в отдельном файле
    
    Returns:
        dict: Информация о созданных файлах
    """
    # Пути по умолчанию
    if input_file is None:
        input_file = os.path.join(os.getcwd(), 'output', 'esklp', 'esklp_full.xlsx')
    
    if output_dir is None:
        output_dir = os.path.join(os.getcwd(), 'output', 'esklp', 'word_reports')
    
    # Создаем директорию для отчетов
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    
    logger.info(f"Загрузка файла {input_file}")
    
    # Проверяем существование файла
    if not os.path.exists(input_file):
        logger.error(f"Файл {input_file} не найден!")
        return {'success': False, 'error': 'File not found'}
    
    try:
        # Загружаем все листы
        xls = pd.ExcelFile(input_file)
        sheet_names = xls.sheet_names
        logger.info(f"Найдены листы: {sheet_names}")
        
        created_files = []
        
        if mode == 'separate':
            # Каждый лист в отдельном документе
            for sheet_name in sheet_names:
                logger.info(f"Обработка листа: {sheet_name}")
                
                # Загружаем данные листа с автоопределением заголовка и фильтрацией Unnamed
                df = load_sheet_with_header_detection(input_file, sheet_name)
                logger.info(f"  Загружено строк: {len(df)}, колонок: {len(df.columns)}")
                
                # Создаем документ Word
                doc = Document()
                
                # Добавляем заголовок
                heading = doc.add_heading(sheet_name, level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Добавляем информацию о дате формирования
                timestamp_para = doc.add_paragraph()
                timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = timestamp_para.add_run(f"Дата формирования: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                run.italic = True
                
                doc.add_paragraph()  # Пустая строка
                
                # Добавляем статистику
                stats_para = doc.add_heading("Статистика", level=2)
                stats_info = [
                    f"Всего записей: {len(df)}",
                    f"Всего колонок: {len(df.columns)}",
                    f"Полностью заполненных колонок: {sum(1 for col in df.columns if df[col].notna().all())}",
                    f"Колонок с пропусками: {sum(1 for col in df.columns if df[col].isna().any())}"
                ]
                for stat in stats_info:
                    doc.add_paragraph(stat, style='List Bullet')
                
                doc.add_paragraph()  # Пустая строка
                
                # Таблица с данными (ограничиваем количество строк для Word)
                max_rows = 1000  # Ограничение для производительности Word
                display_df = df.head(max_rows)
                
                heading_data = doc.add_heading(f"Данные (выборка: {len(display_df)} записей)", level=2)
                
                # Для больших таблиц создаем упрощенный вариант - только первые колонки
                max_cols_to_display = min(15, len(display_df.columns))  # Показываем не более 15 колонок
                
                if len(display_df.columns) > max_cols_to_display:
                    info_para = doc.add_paragraph()
                    info_run = info_para.add_run(
                        f"ℹ️ Показаны первые {max_cols_to_display} колонок из {len(display_df.columns)}. "
                    )
                    info_run.italic = True
                
                # Создаем таблицу
                table = doc.add_table(rows=1, cols=max_cols_to_display)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Заполняем заголовки
                header_row = table.rows[0]
                for i in range(max_cols_to_display):
                    col_name = display_df.columns[i]
                    cell = header_row.cells[i]
                    cell.text = str(col_name)[:25]  # Ограничиваем длину
                    # Делаем текст жирным
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(8)
                    # Устанавливаем фон заголовка
                    set_cell_background(cell, '4472C4')  # Синий цвет
                
                # Заполняем данными
                for _, row in display_df.iterrows():
                    row_cells = table.add_row().cells
                    for i in range(max_cols_to_display):
                        value = row.iloc[i]
                        cell = row_cells[i]
                        cell.text = str(value)[:50] if pd.notna(value) else ''  # Ограничиваем длину
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(7)
                
                doc.add_paragraph()  # Пустая строка
                
                # Добавляем информацию об ограничении
                if len(df) > max_rows or len(df.columns) > max_cols_to_display:
                    warning_para = doc.add_paragraph()
                    parts = []
                    if len(df) > max_rows:
                        parts.append(f"Показаны только первые {max_rows} записей из {len(df)}")
                    if len(df.columns) > max_cols_to_display:
                        parts.append(f"показаны первые {max_cols_to_display} колонок из {len(df.columns)}")
                    warning_run = warning_para.add_run(
                        f"⚠️ {'; '.join(parts)}. Полные данные доступны в исходном Excel файле."
                    )
                    warning_run.italic = True
                
                # Сохраняем документ
                safe_sheet_name = "".join(c for c in sheet_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
                output_file = os.path.join(output_dir, f"{safe_sheet_name}.docx")
                doc.save(output_file)
                logger.info(f"  Сохранено в: {output_file}")
                created_files.append(output_file)
            
            logger.info(f"Создано {len(created_files)} файлов Word")
            
        elif mode == 'single':
            # Все листы в одном документе
            doc = Document()
            
            # Главный заголовок
            main_heading = doc.add_heading("Отчет по справочникам лекарственных препаратов", level=1)
            main_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Дата формирования
            timestamp_para = doc.add_paragraph()
            timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = timestamp_para.add_run(f"Дата формирования: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            run.italic = True
            
            doc.add_page_break()
            
            # Загружаем все листы заранее для анализа взаимосвязей
            all_sheets_data = {}
            all_sheets_pk_fk_info = {}  # Информация о PK/FK для каждого листа
            for sheet_name in sheet_names:
                df = load_sheet_with_header_detection(input_file, sheet_name)
                all_sheets_data[sheet_name] = df
                
                # Анализируем PK/FK для каждой колонки
                pk_fk_info = {}
                for col in df.columns:
                    key_type = detect_pk_fk(df, col)
                    if key_type:
                        pk_fk_info[col] = key_type
                all_sheets_pk_fk_info[sheet_name] = pk_fk_info
            
            # ОБЩАЯ СТАТИСТИКА ПО ВСЕМ ЛИСТАМ
            overall_stats_heading = doc.add_heading("Общая статистика по файлу", level=1)
            
            total_records = sum(len(df) for df in all_sheets_data.values())
            total_columns = sum(len(df.columns) for df in all_sheets_data.values())
            
            overall_stats = [
                f"Всего листов: {len(sheet_names)}",
                f"Суммарное количество записей во всех листах: {total_records:,}",
                f"Суммарное количество колонок во всех листах: {total_columns}",
                f"Средний размер листа: {round(total_records / len(sheet_names), 1):,} записей"
            ]
            
            doc.add_heading("Основные показатели", level=2)
            for stat in overall_stats:
                doc.add_paragraph(stat, style='List Bullet')
            
            # Таблица со статистикой по каждому листу
            doc.add_heading("Статистика по листам", level=2)
            sheet_stats_table = doc.add_table(rows=1, cols=5)
            sheet_stats_table.style = 'Table Grid'
            
            sheet_stats_headers = ['Лист', 'Записей', 'Колонок', 'Полностью заполненных', '% заполненных колонок']
            header_row = sheet_stats_table.rows[0]
            for i, header in enumerate(sheet_stats_headers):
                cell = header_row.cells[i]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            for sheet_name, df in all_sheets_data.items():
                row_cells = sheet_stats_table.add_row().cells
                fully_filled = sum(1 for col in df.columns if df[col].notna().all())
                fill_percent = round(fully_filled / len(df.columns) * 100, 1) if len(df.columns) > 0 else 0
                
                row_cells[0].text = str(sheet_name)[:30]
                row_cells[1].text = f"{len(df):,}"
                row_cells[2].text = str(len(df.columns))
                row_cells[3].text = str(fully_filled)
                row_cells[4].text = f"{fill_percent}%"
            
            doc.add_paragraph()
            
            # АНАЛИЗ ВЗАИМОСВЯЗЕЙ ЛИСТОВ С PK/FK
            doc.add_heading("Анализ взаимосвязей листов", level=2)
            
            # Анализ найденных PK и FK
            all_pks = {}  # {column_name: [(sheet, col), ...]}
            all_fks = {}  # {column_name: [(sheet, col), ...]}
            
            for sheet_name, pk_fk_info in all_sheets_pk_fk_info.items():
                for col, key_type in pk_fk_info.items():
                    if key_type == 'PK':
                        if col not in all_pks:
                            all_pks[col] = []
                        all_pks[col].append((sheet_name, col))
                    elif key_type in ['FK', 'Potential FK']:
                        if col not in all_fks:
                            all_fks[col] = []
                        all_fks[col].append((sheet_name, col))
            
            # Показываем найденные первичные ключи
            if all_pks:
                doc.add_paragraph(f"Найдено первичных ключей (PK): {len(all_pks)}", style='List Bullet')
                pk_table = doc.add_table(rows=1, cols=3)
                pk_table.style = 'Table Grid'
                pk_headers = ['Колонка', 'Лист', 'Уникальность']
                header_row = pk_table.rows[0]
                for i, header in enumerate(pk_headers):
                    cell = header_row.cells[i]
                    cell.text = header
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                for col, locations in sorted(all_pks.items()):
                    for sheet, col_name in locations:
                        df = all_sheets_data[sheet]
                        uniqueness = analyze_column_uniqueness(df, col_name)
                        row_cells = pk_table.add_row().cells
                        row_cells[0].text = str(col)[:40]
                        row_cells[1].text = str(sheet)[:30]
                        row_cells[2].text = f"{uniqueness['uniqueness_percent']}% ({uniqueness['unique']}/{uniqueness['total']})"
                doc.add_paragraph()
            
            # Показываем найденные внешние ключи
            if all_fks:
                doc.add_paragraph(f"Найдено внешних ключей (FK): {len(all_fks)}", style='List Bullet')
                fk_table = doc.add_table(rows=1, cols=3)
                fk_table.style = 'Table Grid'
                fk_headers = ['Колонка', 'Лист', 'Уникальность']
                header_row = fk_table.rows[0]
                for i, header in enumerate(fk_headers):
                    cell = header_row.cells[i]
                    cell.text = header
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                for col, locations in sorted(all_fks.items()):
                    for sheet, col_name in locations:
                        df = all_sheets_data[sheet]
                        uniqueness = analyze_column_uniqueness(df, col_name)
                        row_cells = fk_table.add_row().cells
                        row_cells[0].text = str(col)[:40]
                        row_cells[1].text = str(sheet)[:30]
                        row_cells[2].text = f"{uniqueness['uniqueness_percent']}% ({uniqueness['unique']}/{uniqueness['total']})"
                doc.add_paragraph()
            
            # Анализ потенциальных связей между листами
            doc.add_heading("Потенциальные связи между листами", level=3)
            
            # Ищем колонки с одинаковыми именами, где в одном листе это PK, а в другом FK
            potential_relations = []
            for pk_col, pk_locations in all_pks.items():
                if pk_col in all_fks:
                    fk_locations = all_fks[pk_col]
                    for pk_sheet, _ in pk_locations:
                        for fk_sheet, _ in fk_locations:
                            if pk_sheet != fk_sheet:
                                potential_relations.append({
                                    'from_sheet': fk_sheet,
                                    'to_sheet': pk_sheet,
                                    'column': pk_col,
                                    'type': 'FK → PK'
                                })
            
            if potential_relations:
                doc.add_paragraph(f"Найдено {len(potential_relations)} потенциальных связей:", style='List Bullet')
                rel_table = doc.add_table(rows=1, cols=3)
                rel_table.style = 'Table Grid'
                rel_headers = ['От листа', 'К листу', 'По полю']
                header_row = rel_table.rows[0]
                for i, header in enumerate(rel_headers):
                    cell = header_row.cells[i]
                    cell.text = header
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                for rel in potential_relations[:20]:  # Ограничиваем количество
                    row_cells = rel_table.add_row().cells
                    row_cells[0].text = str(rel['from_sheet'])[:30]
                    row_cells[1].text = str(rel['to_sheet'])[:30]
                    row_cells[2].text = str(rel['column'])[:40]
                
                if len(potential_relations) > 20:
                    doc.add_paragraph(f"... и еще {len(potential_relations) - 20} связей")
            else:
                doc.add_paragraph("Явные связи PK→FK между листами не найдены. Проверьте общие колонки ниже.", style='List Bullet')
            
            # Анализируем потенциальные ключевые поля для связей
            common_column_patterns = ['id', 'код', 'guid', 'uuid', 'name', 'наименование', 'parent', 'родитель']
            
            doc.add_paragraph("\nПоиск общих колонок между листами:", style='List Bullet')
            
            # Собираем все колонки из всех листов
            all_columns = {}
            for sheet_name, df in all_sheets_data.items():
                for col in df.columns:
                    col_lower = str(col).lower()
                    if col_lower not in all_columns:
                        all_columns[col_lower] = []
                    all_columns[col_lower].append(sheet_name)
            
            # Находим колонки, которые встречаются в нескольких листах
            shared_columns = {col: sheets for col, sheets in all_columns.items() if len(sheets) > 1}
            
            if shared_columns:
                doc.add_paragraph(f"Найдено {len(shared_columns)} колонок, встречающихся в нескольких листах:", style='List Bullet')
                
                # Показываем топ-10 наиболее распространенных
                sorted_shared = sorted(shared_columns.items(), key=lambda x: len(x[1]), reverse=True)[:10]
                
                for col_name, sheets in sorted_shared:
                    sheets_str = ', '.join(sheets[:5])
                    if len(sheets) > 5:
                        sheets_str += f" и еще {len(sheets) - 5}"
                    doc.add_paragraph(f"• '{col_name}' → {sheets_str}", style='List Bullet')
            else:
                doc.add_paragraph("Общие колонки между листами не найдены или имена колонок уникальны для каждого листа.")
            
            # Анализ потенциальных внешних ключей
            doc.add_paragraph("\nПотенциальные связи по ключевым полям:", style='List Bullet')
            
            potential_keys = {}
            for sheet_name, df in all_sheets_data.items():
                for col in df.columns:
                    col_lower = str(col).lower()
                    for pattern in common_column_patterns:
                        if pattern in col_lower:
                            if pattern not in potential_keys:
                                potential_keys[pattern] = []
                            potential_keys[pattern].append((sheet_name, col))
            
            if potential_keys:
                for pattern, locations in potential_keys.items():
                    if len(locations) > 1:
                        loc_str = ', '.join([f"{sheet}({col})" for sheet, col in locations[:5]])
                        if len(locations) > 5:
                            loc_str += f" и еще {len(locations) - 5}"
                        doc.add_paragraph(f"• Поле типа '{pattern}': {loc_str}", style='List Bullet')
            
            doc.add_page_break()
            
            # Оглавление
            toc_heading = doc.add_heading("Оглавление", level=2)
            for i, sheet_name in enumerate(sheet_names, 1):
                toc_para = doc.add_paragraph()
                toc_run = toc_para.add_run(f"{i}. {sheet_name}")
                toc_run.bold = True
            
            doc.add_page_break()
            
            # Обработка каждого листа
            for idx, sheet_name in enumerate(sheet_names, 1):
                logger.info(f"Обработка листа: {sheet_name}")
                
                # Заголовок раздела
                section_heading = doc.add_heading(f"{idx}. {sheet_name}", level=1)
                section_heading.page_break_before = True
                
                # Используем предзагруженные данные
                df = all_sheets_data[sheet_name]
                logger.info(f"  Загружено строк: {len(df)}, колонок: {len(df.columns)}")
                
                # Статистика
                stats_heading = doc.add_heading("Статистика", level=2)
                stats_info = [
                    f"Всего записей: {len(df):,}",
                    f"Всего колонок: {len(df.columns)}",
                    f"Полностью заполненных колонок: {sum(1 for col in df.columns if df[col].notna().all())}",
                    f"Колонок с пропусками: {sum(1 for col in df.columns if df[col].isna().any())}"
                ]
                for stat in stats_info:
                    doc.add_paragraph(stat, style='List Bullet')
                
                # Детальная статистика по ВСЕМ колонкам с примером заполнения и PK/FK
                stats_detail_heading = doc.add_heading("Заполненность по колонкам", level=3)
                stats_table = doc.add_table(rows=1, cols=7)
                stats_table.style = 'Table Grid'
                
                # Заголовки таблицы статистики
                stats_headers = ['Колонка', 'Тип ключа', 'Заполнено', 'Пусто', '% заполнения', '% уникальности', 'Пример значения']
                header_row = stats_table.rows[0]
                for i, header in enumerate(stats_headers):
                    cell = header_row.cells[i]
                    cell.text = header
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Данные статистики по ВСЕМ колонкам
                for col in df.columns:
                    non_null_count = df[col].notna().sum()
                    null_count = len(df) - non_null_count
                    fill_percent = round(non_null_count / len(df) * 100, 2) if len(df) > 0 else 0
                    
                    # Анализ уникальности
                    uniqueness = analyze_column_uniqueness(df, col)
                    
                    # Определение PK/FK
                    key_type = detect_pk_fk(df, col)
                    key_display = key_type if key_type else '-'
                    
                    # Получаем пример первого непустого значения
                    example_value = ''
                    first_non_null = df[col][df[col].notna()]
                    if len(first_non_null) > 0:
                        example_value = str(first_non_null.iloc[0])[:50]  # Ограничиваем длину
                    
                    row_cells = stats_table.add_row().cells
                    row_cells[0].text = str(col)[:50]  # Ограничиваем длину названия
                    row_cells[1].text = key_display
                    row_cells[2].text = str(non_null_count)
                    row_cells[3].text = str(null_count)
                    row_cells[4].text = f"{fill_percent}%"
                    row_cells[5].text = f"{uniqueness['uniqueness_percent']}%"
                    row_cells[6].text = example_value
                
                doc.add_paragraph()  # Пустая строка
            
            # Сохраняем единый документ
            output_file = os.path.join(output_dir, "esklp_full_report.docx")
            doc.save(output_file)
            logger.info(f"Сохранено в: {output_file}")
            created_files.append(output_file)
        
        return {
            'success': True,
            'files': created_files,
            'mode': mode,
            'sheets_processed': len(sheet_names)
        }
        
    except Exception as e:
        logger.error(f"Ошибка при создании отчетов: {type(e).__name__}: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return {'success': False, 'error': str(e)}


if __name__ == '__main__':
    import argparse
    
    parser = argparse.ArgumentParser(description='Преобразование Excel отчетов в Word')
    parser.add_argument('--input', '-i', help='Путь к входному файлу Excel')
    parser.add_argument('--output', '-o', help='Директория для выходных файлов Word')
    parser.add_argument('--mode', '-m', choices=['single', 'separate'], default='single',
                       help='Режим: single - все в одном файле, separate - каждый лист отдельно')
    
    args = parser.parse_args()
    
    result = create_word_report_from_excel(
        input_file=args.input,
        output_dir=args.output,
        mode=args.mode
    )
    
    if result['success']:
        print(f"\n✓ Успешно создано {len(result['files'])} файлов Word")
        for f in result['files']:
            print(f"  - {f}")
    else:
        print(f"\n✗ Ошибка: {result.get('error', 'Неизвестная ошибка')}")
        sys.exit(1)
